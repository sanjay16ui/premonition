"""Document ingestion pipeline."""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any

from premonition.copilot.rag.chunking import DocumentChunker
from premonition.copilot.rag.embeddings import EmbeddingService
from premonition.copilot.rag.knowledge import KnowledgeBaseManager
from premonition.copilot.rag.vector_store import VectorDocument, VectorStore


class DocumentIngestionPipeline:
    """Ingest PDF, DOCX, TXT, Markdown into knowledge base + vector store."""

    SUPPORTED_TYPES = {"text", "txt", "md", "markdown", "pdf", "docx", "protocol", "sop"}

    def __init__(
        self,
        knowledge: KnowledgeBaseManager,
        vector_store: VectorStore,
        embedder: EmbeddingService | None = None,
        chunker: DocumentChunker | None = None,
    ) -> None:
        self.knowledge = knowledge
        self.vector_store = vector_store
        self.embedder = embedder or EmbeddingService()
        self.chunker = chunker or DocumentChunker()

    def ingest_text(self, title: str, content: str, doc_type: str = "text", metadata: dict | None = None) -> dict[str, Any]:
        chunks = self.chunker.chunk(content)
        doc = self.knowledge.add(title, content, doc_type, metadata, chunk_count=len(chunks))
        vectors = []
        for chunk in chunks:
            emb = self.embedder.embed(chunk.text)
            vectors.append(VectorDocument(
                id=f"{doc.id}_{chunk.index}",
                text=chunk.text,
                embedding=emb,
                metadata={"title": title, "doc_id": doc.id, "doc_type": doc_type, **(metadata or {})},
                chunk_index=chunk.index,
            ))
        self.vector_store.add(vectors)
        return {"document_id": doc.id, "title": title, "chunks": len(chunks), "doc_type": doc_type}

    def ingest_file(self, path: Path, title: str | None = None, doc_type: str | None = None) -> dict[str, Any]:
        suffix = path.suffix.lower().lstrip(".")
        dtype = doc_type or suffix
        content = self._extract_content(path, suffix)
        return self.ingest_text(title or path.stem, content, dtype, {"source_path": str(path)})

    def _extract_content(self, path: Path, suffix: str) -> str:
        if suffix in {"txt", "md", "markdown"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == "pdf":
            return self._extract_pdf(path)
        if suffix == "docx":
            return self._extract_docx(path)
        return path.read_text(encoding="utf-8", errors="ignore")

    def _extract_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            return f"[PDF content from {path.name} — install pypdf for full extraction]"

    def _extract_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return f"[DOCX content from {path.name} — install python-docx for full extraction]"
